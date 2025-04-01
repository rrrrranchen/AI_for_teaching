import os
from PIL import Image
from pptx import Presentation
from docx import Document
import cv2
from mutagen.mp4 import MP4
from mutagen.mp3 import MP3
import PyPDF2
from datetime import datetime
from typing import Dict, Optional

class FileParser:
    """多媒体文件元数据解析工具类"""
    
    @staticmethod
    def parse_image(filepath: str) -> Dict:
        """解析图片文件(JPEG/PNG/GIF等)
        
        Args:
            filepath: 图片文件路径
            
        Returns:
            包含以下键的字典:
            - resolution (str): 分辨率如 '1920x1080'
            - format (str): 文件格式如 'JPEG'
            - color_mode (str): 颜色模式如 'RGB'
            - file_size (int): 文件大小(KB)
        """
        with Image.open(filepath) as img:
            return {
                'resolution': f"{img.width}x{img.height}",
                'format': img.format,
                'color_mode': img.mode,
                'file_size': os.path.getsize(filepath) // 1024  # KB
            }
    
    @staticmethod
    def parse_video(filepath: str) -> Dict:
        """解析视频文件(MP4/MOV等)
        
        Args:
            filepath: 视频文件路径
            
        Returns:
            包含以下键的字典:
            - resolution (str): 分辨率如 '1920x1080'
            - duration (int): 视频时长(秒)
            - frame_rate (float): 帧率
            - format (str): 文件格式如 'mp4'
            - file_size (int): 文件大小(MB)
        """
        # 使用OpenCV获取基础信息
        cap = cv2.VideoCapture(filepath)
        width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = frame_count / fps if fps > 0 else 0
        cap.release()
        
        # 使用mutagen获取更多元数据
        try:
            video = MP4(filepath)
            meta_duration = video.info.length
        except:
            meta_duration = duration
            
        return {
            'resolution': f"{width}x{height}",
            'duration': int(meta_duration),
            'frame_rate': fps,
            'format': 'mp4',
            'file_size': os.path.getsize(filepath) // (1024 * 1024)  # MB
        }
    
    @staticmethod
    def parse_ppt(filepath: str) -> Dict:
        """解析PPT文件(PPTX)
        
        Args:
            filepath: PPT文件路径
            
        Returns:
            包含以下键的字典:
            - slide_count (int): 幻灯片页数
            - author (str): 作者
            - created_date (datetime): 创建时间
            - format (str): 文件格式如 'pptx'
            - file_size (int): 文件大小(KB)
        """
        prs = Presentation(filepath)
        return {
            'slide_count': len(prs.slides),
            'author': prs.core_properties.author or '',
            'created_date': prs.core_properties.created or datetime.now(),
            'format': 'pptx',
            'file_size': os.path.getsize(filepath) // 1024  # KB
        }
    
    @staticmethod
    def parse_doc(filepath: str) -> Dict:
        """解析Word文档(DOCX)
        
        Args:
            filepath: Word文件路径
            
        Returns:
            包含以下键的字典:
            - page_count (int): 估算页数
            - word_count (int): 估算字数
            - author (str): 作者
            - created_date (datetime): 创建时间
            - format (str): 文件格式如 'docx'
            - file_size (int): 文件大小(KB)
        """
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs]
        text_content = '\n'.join(paragraphs)
        
        return {
            'page_count': len(doc.paragraphs) // 50 + 1,  # 估算页数
            'word_count': len(text_content.split()),
            'author': doc.core_properties.author or '',
            'created_date': doc.core_properties.created or datetime.now(),
            'format': 'docx',
            'file_size': os.path.getsize(filepath) // 1024  # KB
        }
    
    @staticmethod
    def parse_audio(filepath: str) -> Dict:
        """解析音频文件(MP3)
        
        Args:
            filepath: 音频文件路径
            
        Returns:
            包含以下键的字典:
            - duration (int): 时长(秒)
            - bitrate (int): 比特率
            - format (str): 文件格式如 'mp3'
            - file_size (int): 文件大小(KB)
        """
        audio = MP3(filepath)
        return {
            'duration': int(audio.info.length),
            'bitrate': audio.info.bitrate,
            'format': 'mp3',
            'file_size': os.path.getsize(filepath) // 1024  # KB
        }

    @staticmethod
    def parse_pdf(filepath: str) -> Dict:
        """解析PDF文件
        
        Args:
            filepath: PDF文件路径
            
        Returns:
            包含以下键的字典:
            - page_count (int): 总页数
            - word_count (int): 估算字数
            - author (str): 作者
            - title (str): 文档标题
            - created_date (datetime): 创建时间
            - format (str): 文件格式如 'pdf'
            - file_size (int): 文件大小(KB)
        """
        meta = {
            'format': 'pdf',
            'file_size': os.path.getsize(filepath) // 1024,  # KB
            'page_count': 0,
            'word_count': 0,
            'author': '',
            'title': '',
            'created_date': None
        }
        
        try:
            with open(filepath, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                
                meta['page_count'] = len(pdf.pages)
                
                # 提取文本前1000字符用于分析
                sample_text = ""
                for page in pdf.pages[:2]:  # 只读取前两页节省资源
                    text = page.extract_text()
                    if text:
                        sample_text += text[:500]
                meta['word_count'] = len(sample_text.split())
                
                # 提取元数据
                if hasattr(pdf, 'metadata') and pdf.metadata:
                    meta.update({
                        'author': pdf.metadata.get('/Author', ''),
                        'title': pdf.metadata.get('/Title', ''),
                        'creator': pdf.metadata.get('/Creator', ''),
                        'producer': pdf.metadata.get('/Producer', ''),
                        'subject': pdf.metadata.get('/Subject', ''),
                        'keywords': pdf.metadata.get('/Keywords', '')
                    })
                    
                    # 解析PDF日期格式 (D:YYYYMMDDHHmmSS)
                    if '/CreationDate' in pdf.metadata:
                        date_str = pdf.metadata['/CreationDate'][2:16]  # D:20230821195643
                        try:
                            meta['created_date'] = datetime.strptime(date_str, '%Y%m%d%H%M%S')
                        except ValueError:
                            pass
                            
        except Exception as e:
            meta['error'] = f"PDF解析错误: {str(e)}"
        
        return meta

    @classmethod
    def parse_file(cls, filepath: str, file_type: Optional[str] = None) -> Dict:
        """自动解析文件元数据
        
        Args:
            filepath: 文件路径
            file_type: 可选的文件类型(如'pdf','mp4')
            
        Returns:
            包含元数据的字典，出错时包含error键
        """
        if not file_type:
            ext = os.path.splitext(filepath)[1].lower()[1:]
            file_type = ext if ext in [
                'jpg', 'jpeg', 'png', 'gif',  # 图片
                'mp4', 'mov', 'avi',          # 视频
                'pptx', 'ppt',                # PPT
                'docx', 'doc',                # Word
                'mp3', 'wav',                 # 音频
                'pdf'                         # PDF
            ] else None
        
        try:
            if file_type in ['jpg', 'jpeg', 'png', 'gif']:
                return cls.parse_image(filepath)
            elif file_type in ['mp4', 'mov', 'avi']:
                return cls.parse_video(filepath)
            elif file_type in ['pptx', 'ppt']:
                return cls.parse_ppt(filepath)
            elif file_type in ['docx', 'doc']:
                return cls.parse_doc(filepath)
            elif file_type in ['mp3', 'wav']:
                return cls.parse_audio(filepath)
            elif file_type == 'pdf':
                return cls.parse_pdf(filepath)
            else:
                return {'error': '不支持的文件类型'}
        except Exception as e:
            return {'error': f'文件解析失败: {str(e)}'}