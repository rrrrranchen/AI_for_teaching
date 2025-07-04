import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.shared import Inches
import os
from datetime import datetime
import random
import string

# 配置参数
num_questions = 10  # 题目数量
num_students = 50   # 学生数量
question_types = ['填空', '选择', '简答']  # 题目类型
choices = ['A', 'B', 'C', 'D']  # 选择题的选项

# 生成模拟数据
def generate_mock_data(num_questions, num_students, question_types, choices):
    data = {
        'q_id': [f'q{i + 1}' for i in range(num_questions)],
        'type': [random.choice(question_types) for _ in range(num_questions)],
        'answers': []
    }

    scores = {
        'q_id': [f'q{i + 1}' for i in range(num_questions)]
    }

    # 为每个学生生成作答记录和得分
    for i in range(num_students):
        student_id = f's{i + 1}'
        student_answers = []
        student_scores = []

        for j in range(num_questions):
            q_type = data['type'][j]

            # 根据题目类型生成作答记录和得分
            if q_type == '选择':
                answer = random.choice(choices)
                correct_answer = random.choice(choices)
                score = 1.0 if answer == correct_answer else 0.0
            elif q_type == '填空':
                answer = random.randint(0, 1)
                correct_answer = random.randint(0, 1)
                score = 1.0 if answer == correct_answer else 0.0
            elif q_type == '简答':
                answer = ''.join(random.choices(string.ascii_letters, k=2))
                correct_answer = ''.join(random.choices(string.ascii_letters, k=2))
                score = round(random.uniform(0, 1), 2)  # 简答题得分随机

            student_answers.append(answer)
            student_scores.append(score)

        data[student_id] = student_answers
        scores[student_id] = student_scores

        # 保存正确答案
        data['answers'].append(correct_answer)

    return data, scores





# 中文显示配置
plt.rcParams['font.sans-serif'] = ['SimHei']
plt.rcParams['axes.unicode_minus'] = False

project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
WORDTEMPLATE_FOLDER = os.path.join(project_root, 'static', 'template')
REPORT_FOLDER=os.path.join(project_root,'static','uploads','reports')
PNG_FOLDER=os.path.join(project_root,'static','uploads','pngs')
filename='template.docx'
templatepath=os.path.join(WORDTEMPLATE_FOLDER, filename)

def analyze_data():
    """
    data = {
        'q_id': ['q1', 'q2'],
        's1': [1, 'A'],
        's2': [0, 'B'],
        'type': ['填空', '选择'],
        'answers': [0, 'C']
    }
    scores = {
        'q_id': ['q1', 'q2'],
        's1': [0.0, 0.0],
        's2': [1.0, 0.0],
    }
    """
    pass


def bar_chart(scores, output_dir):
    plt.figure(figsize=(12, 6))
    df = pd.DataFrame({
        'Question': scores['q_id'],
        'ScoreRate': np.mean([v for k, v in scores.items() if k != 'q_id'], axis=0)
    })

    ax = sns.barplot(
        x='Question',
        y='ScoreRate',
        data=df,
        palette=['#ff0000' if x < 0.6 else '#00ff00' for x in df['ScoreRate']],
        edgecolor='black',
        linewidth=1.5
    )

    plt.axhline(0.6, color='r', linestyle='--', alpha=0.3)
    plt.title("各题目平均得分", pad=20)
    plt.xlabel("题目编号")
    plt.ylabel("得分率")
    plt.ylim(0, 1)

    for p in ax.patches:
        ax.annotate(
            f"{p.get_height():.1%}",
            (p.get_x() + p.get_width() / 2., p.get_height()),
            ha='center',
            va='center',
            xytext=(0, 10),
            textcoords='offset points',
            fontsize=12
        )

    # 保存图像
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, "平均得分.png")
    plt.savefig(path, bbox_inches='tight', dpi=300)
    plt.close()

    return path


def question_charts(data, scores, output_dir):
    num = len(data['q_id'])
    paths = [''] * num

    for i in range(num):
        q_id = data['q_id'][i]
        q_type = data['type'][i]

        plt.figure(figsize=(8, 6))

        # ================== 选择题处理 ==================
        if q_type == '选择':
            labels = ['A', 'B', 'C', 'D']
            records = [l for k, l in data.items()
                       if k not in ['q_id', 'type', 'answers']]
            record = [s[i] for s in records]
            sizes = [max(record.count(k) / len(record), 1e-5) for k in labels]
            colors = ['#00ff00' if k == data['answers'][i] else '#ff0000'
                      for k in labels]

            plt.pie(sizes, labels=labels, autopct='%1.1f%%', colors=colors)
            plt.title(f"题目{q_id}选项分布")

        # ================== 简答题处理 ================== 分数为一区间值
        elif q_type == '简答':
            q_scores = [s[i] for s in scores.values() if s is not scores['q_id']]

            # 分箱统计
            bins = [
                ('0分', lambda x: x == 0),
                ('低分', lambda x: 0 < x < 0.6),
                ('及格', lambda x: 0.6 <= x < 0.8),
                ('高分', lambda x: 0.8 <= x < 1.0),
                ('满分', lambda x: x == 1)
            ]
            sizes = [max(len(list(filter(func, q_scores))) / len(q_scores), 1e-5)
                     for (label, func) in bins]
            labels = [label for (label, func) in bins]
            colors = ['#ff0000', '#990099', '#000099', '#009999', '#00ff00']

            plt.pie(sizes, labels=labels, autopct='%1.1f%%', colors=colors)
            plt.title(f"题目{q_id}得分分布")

        # ================== 其他题型处理 ================== 分数为1/0
        else:
            q_scores = [s[i] for s in scores.values() if s is not scores['q_id']]
            count_1 = int(sum(q_scores))
            count_0 = len(q_scores) - count_1
            total = len(q_scores)
            percent_1 = count_1 / total
            percent_0 = count_0 / total

            df = pd.DataFrame({
                "Category": ["正确", "错误"],
                "Percentage": [percent_1, percent_0]
            })

            plt.figure(figsize=(8, 3))
            sns.barplot(
                x="Percentage",
                y="Category",
                data=df,
                palette=["#ff0000", "#00ff00"],
                errorbar=None  # 不显示误差条
            )

            plt.title(f"题目{q_id}正确率分布")
            plt.xlabel("比例")
            plt.ylabel("")
            plt.xlim(0, 1)

            # 添加百分比标签
            for index, value in enumerate(df["Percentage"]):
                plt.text(
                    value,
                    index,
                    f"{value:.1%}",
                    va="center",
                    ha="left",
                    fontsize=12
                )
            # 多个分数值
            # q_scores = [s[i] for s in scores.values() if s is not scores['q_id']]
            #
            # unique_scores = sorted(list(set(q_scores)))
            # counts = [q_scores.count(s) for s in unique_scores]
            # total = len(q_scores)
            # percentages = [c / total for c in counts]
            #
            # # 横向柱状图
            # plt.barh(unique_scores, percentages, color='#0099ff')
            # plt.title(f"题目{q_id}得分分布")
            # plt.xlabel("比例")
            # plt.ylabel("得分")
            # plt.xlim(0, 1)
            #
            # for idx, (score, pct) in enumerate(zip(unique_scores, percentages)):
            #     plt.text(pct, idx, f"{pct:.1%}",
            #              va='center', ha='left', fontsize=10)

        # ================== 保存文件 ==================
        os.makedirs(output_dir, exist_ok=True)
        filename = f"{q_id}{q_type}分析.png"
        path = os.path.join(output_dir, filename)
        plt.savefig(path, bbox_inches='tight', dpi=300)
        plt.close()
        paths[i] = path

    return paths


def set_word_font(doc, font_name):
    for p in doc.paragraphs:
        if p.text:
            # 设置段落字体
            run = p.runs[0]
            run.font.name = font_name
            # run.font.size = Pt(12)  # 设置字体大小


def generate_report(data, scores):
    output_dir = PNG_FOLDER
    # os.makedirs(output_dir, exist_ok=True)
    file_name = os.path.join(REPORT_FOLDER, f"学生作答分析报告{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

    # 1. 生成总得分柱状图
    bar_path = bar_chart(scores, output_dir)

    # 2. 生成每个题目的饼图
    pie_paths = question_charts(data, scores, output_dir)

    # 3. 生成报告
    doc = Document(templatepath)
    # title = doc.add_heading('学生作答综合分析报告', 0)
    # title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 总得分图
    # doc.add_heading('各题目得分率对比', level=1)
    # doc.add_picture(bar_path, width=Inches(6))
    for p in doc.paragraphs:
        if p.text == '平均得分.png':
            p.clear()
            run = p.add_run()
            run.add_picture(bar_path, width=Inches(6))

    # 题目分析
    # doc.add_heading('各题目详细分析', level=1)
    for idx, path in enumerate(pie_paths):
        if data['type'][idx] == '选择':
            records = [l for k, l in scores.items()
                       if k != 'q_id']
            record = [s[idx] for s in records]
            # doc.add_heading(f"题目{idx+1}选项分布\n"
            #                 f"正确答案: {data['answers'][idx]}\n"
            #                 f"正确率: {100 * round(record.count(1.0) / len(record), 4)}%", level=2)
            p = doc.add_paragraph()
            run = p.add_run(f"题目{idx + 1}选项分布\n"
                            f"正确答案: {data['answers'][idx]}\n"
                            f"正确率: {100 * round(record.count(1.0) / len(record), 4)}%")
            run.font.color.rgb = RGBColor(0, 0, 0)
            # run.font.name = '楷体'
            run.font.name = "KaiTi"  # 楷体的英文字体名
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "KaiTi")  # 强制设置东亚字体
            run.font.size = Pt(12)
        else:
            records = [l for k, l in scores.items()
                       if k != 'q_id']
            record = [s[idx] for s in records]
            # doc.add_heading(f"题目{idx + 1}答案解析\n"
            #                 f"正确答案: {data['answers'][idx]}\n"
            #                 f"正确率: {100 * round(record.count(1.0) / len(record), 4)}%", level=2)
            p = doc.add_paragraph()
            run = p.add_run(f"题目{idx + 1}答案解析\n"
                            f"正确答案: {data['answers'][idx]}\n"
                            f"正确率: {100 * round(record.count(1.0) / len(record), 4)}%")
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = "KaiTi"  # 楷体的英文字体名
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "KaiTi")  # 强制设置东亚字体
            run.font.size = Pt(12)
        doc.add_picture(path, width=Inches(4))

    # set_word_font(doc, '楷体')
    doc.save(file_name)
    print(f'报告已保存为: {file_name}')



data, scores = generate_mock_data(num_questions, num_students, question_types, choices)
generate_report(data, scores)  # Warning是正常的!

