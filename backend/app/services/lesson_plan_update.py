from openai import OpenAI
import time
from docx import Document
import markdown
import pdfkit

# 设置 API Key 和 DeepSeek API 地址
key = 'sk-b7550aa67ed840ffacb5ca051733802c'
client = OpenAI(api_key=key, base_url="https://api.deepseek.com")


# 逐字打印函数
def printChar(text, delay=0.05):
    for char in text:
        print(char, end='', flush=True)
        time.sleep(delay)
    print()


# 生成预备知识检测题和问卷
def generate_pre_class_questions(course_content):
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {
                "role": "system",
                "content": "你是教学设计专家，请根据教师提供的课程内容生成一份预备知识检测练习题和一份调查问卷，用于了解学生对相关内容的了解程度，同时也要生成习题的答案。"
            },
            {
                "role": "user",
                "content": f"课程内容如下：\n{course_content}"
            }
        ],
        stream=False
    )
    return response.choices[0].message.content

# 将习题和问卷生成word文档

# def save_to_word(content, filename="output.docx"):
#     # 解析 Markdown 为纯文本（去掉 #，但保留结构）
#     md_lines = content.split("\n")
#
#     # 创建 Word 文档
#     doc = Document()
#
#     for line in md_lines:
#         if line.startswith("# "):  # 一级标题
#             doc.add_heading(line[2:], level=1)
#         elif line.startswith("## "):  # 二级标题
#             doc.add_heading(line[3:], level=2)
#         elif line.startswith("### "):  # 三级标题
#             doc.add_heading(line[4:], level=3)
#         else:
#             doc.add_paragraph(line)
#
#     # 保存 Word 文件
#     doc.save(filename)
#     print(f"\n✅ Word 文件已保存为：{filename}")


# 生成结构化教案（按六大模块分段）
def generate_lesson_plans(course_content, student_feedback, student_level):
    """
    根据学生反馈和群体水平，生成一个教学方案
    :param course_content: 课程内容
    :param student_feedback: 学生反馈
    :param student_level: 学生群体水平（如 '良好', '一般', '薄弱'）
    :return: 教案内容
    """
    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{
            "role": "system",
            "content": (
                "你是教学设计专家，请根据教师提供的课程内容和学生答题反馈，"
                f"尽可能详细地设计适用于{student_level}群体的教案。\n\n"
                "每套教案请严格按照以下六个模块详细展开，并满足以下要求：\n\n"
                "1. 教学目标：明确列出3-5个可衡量的学习目标。\n"
                "2. 教学重难点：说明本节课的重点与难点，并简述突破策略。\n"
                "3. 教学内容：详细展开教学内容的结构，包括知识点的层次划分与逻辑关系。\n"
                "4. 教学时间安排：请按总课时45分钟，分配每一阶段的时间（含每个活动的分钟数）。\n"
                "5. 教学过程：\n"
                "    - 请按照“导入、讲授、互动、小结”顺序编排。\n"
                "    - 必须设计不少于三个互动环节，如小组讨论、角色扮演、实时问答、投票、课堂游戏等，同时在生成教学设计方案的时候要明确指出哪部分是互动环节。\n"
                "    - 每个环节说明教学方法、活动安排、教师与学生的行为、时间分配、使用的工具与材料、预期学习成果。\n"
                "6. 课后作业：布置有层次的作业任务，至少包含基础题与拓展题。\n\n"
                "请使用 Markdown 格式输出，便于后续整理归档，但不需要将内容加上markdown注释。每份教案的字数不得少于1000字。"
            )
        },
        {
            "role": "user",
            "content": f"课程内容如下：\n{course_content}\n\n学生反馈如下：\n{student_feedback}"
        }],
        stream=False
    )
    return response.choices[0].message.content



# 推荐指数
def evaluate_recommendation(student_feedback):
    """
    根据学生反馈，评估三套教案的适用性，返回推荐指数。
    推荐指数采用百分制。
    """
    feedback = student_feedback.lower()
    good_keywords = ['掌握好', '熟练', '掌握较好', '轻松', '容易', '良好', '优秀']
    average_keywords = ['一般', '有些困难', '不熟悉', '基本', '部分']
    weak_keywords = ['困难', '不会', '不懂', '模糊', '很难', '不了解']

    good_score = sum(feedback.count(kw) for kw in good_keywords)
    average_score = sum(feedback.count(kw) for kw in average_keywords)
    weak_score = sum(feedback.count(kw) for kw in weak_keywords)

    total_score = good_score + average_score + weak_score

    if total_score == 0:
        return {'掌握良好': 33, '掌握一般': 34, '掌握薄弱': 33}  # 默认均分推荐指数

    recommendation = {
        '掌握良好': int(good_score / total_score * 100),
        '掌握一般': int(average_score / total_score * 100),
        '掌握薄弱': int(weak_score / total_score * 100)
    }
    return recommendation


# 教案保存
def save_to_markdown(content, filename="教案.md"):
    """保存教案为 Markdown 文件"""
    with open(filename, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"\n✅ 教案已保存为 Markdown 文件：{filename}")


# 教案保存为PDF
# def save_to_pdf(content, filename="教案.pdf"):
#     # 将 Markdown 转换为 HTML
#     html_content = markdown.markdown(content)
#
#     # HTML 模板
#     html_template = f"""
#     <html>
#     <head>
#         <meta charset="utf-8">
#         <style>
#             body {{ font-family: "SimSun", serif; line-height: 1.6; margin: 40px; }}
#             h1 {{ text-align: center; color: #2c3e50; }}
#             h2 {{ color: #2c3e50; margin-top: 25px; }}
#             .section {{ margin-bottom: 20px; }}
#         </style>
#     </head>
#     <body>
#         <h1>教案设计</h1>
#         <div class="section">{html_content}</div>
#     </body>
#     </html>
#     """
#
#     # 手动指定 wkhtmltopdf 路径
#     pdfkit_config = pdfkit.configuration(wkhtmltopdf=r"E:\Software\wkhtmltopdf\bin\wkhtmltopdf.exe")
#
#     # 生成 PDF
#     pdfkit.from_string(html_template, filename, configuration=pdfkit_config)
#     print(f"\n✅ PDF 已保存为：{filename}")


# 推荐指数
def evaluate_recommendation(student_feedback):
    """
    根据学生反馈，评估三套教案的适用性，返回推荐指数。
    推荐指数采用百分制。
    """
    feedback = student_feedback.lower()
    good_keywords = ['掌握好', '熟练', '掌握较好', '轻松', '容易', '良好', '优秀']
    average_keywords = ['一般', '有些困难', '不熟悉', '基本', '部分']
    weak_keywords = ['困难', '不会', '不懂', '模糊', '很难', '不了解']

    good_score = sum(feedback.count(kw) for kw in good_keywords)
    average_score = sum(feedback.count(kw) for kw in average_keywords)
    weak_score = sum(feedback.count(kw) for kw in weak_keywords)

    total_score = good_score + average_score + weak_score

    if total_score == 0:
        return {'掌握良好': 33, '掌握一般': 34, '掌握薄弱': 33}  # 默认均分推荐指数

    recommendation = {
        '掌握良好': int(good_score / total_score * 100),
        '掌握一般': int(average_score / total_score * 100),
        '掌握薄弱': int(weak_score / total_score * 100)
    }
    return recommendation


# 主程序
def teacher_assistant():
    print("🎓 欢迎使用【教师备课助手】\n")

    lesson_plan_versions = []
    course_content = input("📚 请输入本节课的教学内容（学科、章节或知识点）：\n")

    print("\n🤖 正在生成预备知识检测题与学生问卷，请稍候...\n")
    questions = generate_pre_class_questions(course_content)
    print("✅ 以下是为本节课自动生成的内容：\n")
    printChar(questions)

    save = input("\n💾 是否将检测习题及问卷保存为 Markdown？(y/n): ")
    if save.lower() == 'y':
        save_to_markdown(questions, f"预备知识检测题和问卷.md")

    student_feedback = input("\n📊 请输入学生答题结果、共性问题或学习反馈（可简述）：\n")

    student_levels = ['掌握良好', '掌握一般', '掌握薄弱']
    recommendation_scores = evaluate_recommendation(student_feedback)
    level_versions = {level: 1 for level in student_levels}

    # 第一次生成三份教案
    for level in student_levels:
        print(f"\n🤖 正在为【{level}】群体生成个性化教案，请稍候...\n")
        lesson_plans = generate_lesson_plans(course_content, student_feedback, level)
        round_num = level_versions[level]
        filename = f"教案_{level}_V{round_num}.md"
        recommendation_score = recommendation_scores.get(level, 0)

        lesson_plan_versions.append({
            'level': level,
            'content': lesson_plans,
            'filename': filename,
            'recommendation': recommendation_score
        })

        print(f"\n📘 {level}群体的教案（推荐指数：{recommendation_score}%）：\n")
        printChar(lesson_plans)

        save = input(f"\n💾 是否将【{level}】群体的教案保存为 Markdown？(y/n): ")
        if save.lower() == 'y':
            save_to_markdown(lesson_plans, filename)

        level_versions[level] += 1

    # 满意度循环
    while True:
        satisfied = input("\n📋 您是否满意这三份教学设计？(y/n): ")
        if satisfied.lower() == 'y':
            break

        # 询问哪个教案不满意
        print("🤔 请问您对哪一个群体的教案不满意？")
        print("选项：掌握良好 / 掌握一般 / 掌握薄弱")
        unsatisfied_level = input("请输入不满意的学生群体：")

        if unsatisfied_level not in student_levels:
            print("⚠️ 输入无效，请重新输入一个有效的群体。")
            continue

        revise_prompt = input(f"\n✏️ 请说明您希望修改【{unsatisfied_level}】教案的内容或方向：\n")

        # 拼接修改意见重新生成
        print(f"\n🤖 正在根据修改意见重新生成【{unsatisfied_level}】群体教案...\n")
        updated_plan = generate_lesson_plans(course_content, student_feedback + "\n教师修改建议：" + revise_prompt, unsatisfied_level)

        version = level_versions[unsatisfied_level]
        new_filename = f"教案_{unsatisfied_level}_V{version}.md"
        new_recommendation = recommendation_scores.get(unsatisfied_level, 0)

        # 替换旧教案，保留所有版本
        lesson_plan_versions.append({
            'level': unsatisfied_level,
            'content': updated_plan,
            'filename': new_filename,
            'recommendation': new_recommendation
        })

        print(f"\n📘 教案已重新生成（版本：V{version}，推荐指数：{new_recommendation}%）：\n")
        printChar(updated_plan)

        save = input(f"\n💾 是否将更新后的【{unsatisfied_level}】教案保存为 Markdown？(y/n): ")
        if save.lower() == 'y':
            save_to_markdown(updated_plan, new_filename)

        level_versions[unsatisfied_level] += 1

    # 输出所有生成的教案
    print("\n📂 所有生成的教学方案文件清单：")
    for plan in lesson_plan_versions:
        preview = plan['content'][:20].replace('\n', ' ') + "..."
        print(f"{plan['level']} | 推荐指数：{plan['recommendation']}% | {plan['filename']}\n→ {preview}")

    print("\n🎉 教学设计完成，祝你上课顺利！")

    return lesson_plan_versions



# 启动程序
if __name__ == '__main__':
    lesson_plans = teacher_assistant()
