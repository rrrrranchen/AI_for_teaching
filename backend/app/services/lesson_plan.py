import json
import re
from venv import logger
from openai import OpenAI
import time
from docx import Document
import markdown
import pdfkit
import requests
from sqlalchemy import and_, false, true
from typing import Any, List, Dict, Optional, Tuple, Union
from concurrent.futures import ThreadPoolExecutor, as_completed
import json
from typing import Dict

from app.utils.ai_chat import _retrieve_chunks_from_multiple_dbs
# 设置 API Key 和 DeepSeek API 地址
key = 'sk-b7550aa67ed840ffacb5ca051733802c'
client = OpenAI(api_key=key, base_url="https://api.deepseek.com")


# 逐字打印函数
def printChar(text, delay=0.05):
    for char in text:
        print(char, end='', flush=True)
        time.sleep(delay)
    print()


# 生成预备知识检测题和问卷
import json



def generate_ai_analysis(knowledge_point_name: str, student_answers: Dict[str, Any]) -> str:
    """
    调用 AI 接口，围绕特定知识点分析学生的学习掌握情况，并返回简短的分析报告。

    参数:
        knowledge_point_name: 知识点名称
        student_answers: JSON 格式的学生作答情况，包含每个题目的学生作答数据

    返回:
        纯文本格式的分析报告，包含：
        - 分析摘要
        - 错误分析
        - 薄弱环节
    """
    system_prompt = """你是教学分析专家，请根据学生对知识点题目的作答情况，生成简短的分析报告。
报告必须包含以下内容：
1. 学生对知识点的掌握情况（整体分析）
2. 常见错误点及错误原因分析
3. 学生理解的难点、薄弱环节（具体知识点内容）

输出格式：
分析摘要：
[在这里输出学生对知识点的掌握情况]

错误分析：
[在这里输出常见错误点及错误原因]

薄弱环节：
[在这里输出学生理解的难点和薄弱环节]
"""

    try:
        # 调用 AI 接口
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": f"知识点名称：{knowledge_point_name}\n学生作答数据：\n{json.dumps(student_answers, ensure_ascii=False)}"
                }
            ],
            temperature=0.5
        )

        # 解析 AI 返回的内容
        ai_response = response.choices[0].message.content

        # 提取分析报告内容
        analysis_summary = re.search(r'分析摘要：(.*?)错误分析', ai_response, re.DOTALL)
        error_analysis = re.search(r'错误分析：(.*?)薄弱环节', ai_response, re.DOTALL)
        weak_points = re.search(r'薄弱环节：(.*?)$', ai_response, re.DOTALL)

        # 构建纯文本格式的分析报告
        analysis_report = f"""
分析摘要：
{analysis_summary.group(1).strip() if analysis_summary else '无法生成摘要'}

错误分析：
{error_analysis.group(1).strip() if error_analysis else '无法生成错误分析'}

薄弱环节：
{weak_points.group(1).strip() if weak_points else '无法检测薄弱环节'}
"""

        return analysis_report

    except Exception as e:
        print(f"AI 分析出错: {str(e)}")
        # 出错时返回默认值
        return """
分析摘要：
AI 分析失败，无法生成摘要

错误分析：
无法生成错误分析

薄弱环节：
无法检测薄弱环节
"""
def generate_pre_class_questions(course_content: str) -> List[Dict[str, Union[str, int]]]:
    """
    生成课前预习题目（完整优化版）
    
    参数:
        course_content: 课程内容文本
        
    返回:
        题目列表，每个题目包含:
        - type: 题目类型 (choice/fill/short_answer)
        - content: 题目内容（选择题为JSON字符串）
        - correct_answer: 正确答案
        - difficulty: 难度等级 (1-5)
        - timing: 题目时间类型
        
    异常处理:
        - 自动修复格式问题
        - 提供默认值保证始终返回有效数据
    """
    system_prompt = """你是教学设计专家，请根据课程内容生成3-5道预备知识检测题。要求：
1. 返回JSON格式，包含questions数组
2. 每个题目包含:
   - type: 题目类型(choice/fill/short_answer)
   - content: 题目内容（选择题需包含question和options字段）
   - correct_answer: 正确答案
   - difficulty: 难度(1-5)
3. 选择题示例格式:
   {
     "type": "choice",
     "content": {
       "question": "问题文本",
       "options": ["A.选项1", "B.选项2"]
     },
     "correct_answer": "A"
   }
4. 用中文回答"""

    try:
        # 调用AI接口
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"课程内容:\n{course_content}"}
            ],
            response_format={"type": "json_object"},
            temperature=0.7
        )
        
        # 解析响应
        ai_data = json.loads(response.choices[0].message.content)
        raw_questions = ai_data.get("questions", [])
        
        processed_questions = []
        for i, q in enumerate(raw_questions, 1):
            question = {
                "type": validate_question_type(q.get("type")),
                "difficulty": clamp_difficulty(q.get("difficulty", 3)),
                "timing": "pre_class"
            }
            
            # 处理题目内容
            if question["type"] == "choice":
                content_data = parse_choice_content(q)
                question["content"] = json.dumps(content_data, ensure_ascii=False)
                question["correct_answer"] = validate_answer(
                    q.get("correct_answer"), 
                    content_data["options"]
                )
            else:
                question["content"] = str(q.get("content", f"问题{i}"))[:1000]
                question["correct_answer"] = str(q.get("correct_answer", ""))[:500]
            
            processed_questions.append(question)
        
        return processed_questions or [get_default_question(course_content)]

    except Exception as e:
        print(f"题目生成失败: {str(e)}")
        return [get_default_question(course_content)]

# 辅助函数 --------------------------------------------------

def validate_question_type(q_type: str) -> str:
    """验证题目类型"""
    valid_types = {"choice", "fill", "short_answer"}
    return q_type.lower() if q_type.lower() in valid_types else "choice"

def clamp_difficulty(diff: Union[int, float]) -> int:
    """限制难度范围1-5"""
    try:
        return max(1, min(5, int(diff)))
    except (TypeError, ValueError):
        return 3

def parse_choice_content(question_data: Dict) -> Dict:
    """解析选择题内容"""
    content = question_data.get("content", {})
    
    # 处理嵌套格式
    if isinstance(content, dict):
        question = content.get("question", "选择题")
        options = content.get("options", [])
    else:
        question = str(content)
        options = question_data.get("options", [])
    
    # 标准化选项
    if not isinstance(options, list):
        options = []
    
    # 确保每个选项有编号
    formatted_options = []
    for i, opt in enumerate(options[:10]):  # 最多10个选项
        opt_str = str(opt)
        if not opt_str.startswith(("A.", "B.", "C.", "D.", "E.", "F.")):
            prefix = chr(65 + i) + "."  # A., B., etc.
            opt_str = f"{prefix} {opt_str}"
        formatted_options.append(opt_str[:200])  # 限制选项长度
    
    # 确保至少2个选项
    if len(formatted_options) < 2:
        formatted_options = ["A. 选项A", "B. 选项B"]
    
    return {
        "question": str(question)[:500],
        "options": formatted_options
    }

def validate_answer(answer: str, options: List[str]) -> str:
    """验证选择题答案"""
    if not options:
        return "A"
    
    # 提取有效选项字母
    valid_choices = [opt[0] for opt in options if len(opt) > 0]
    if not valid_choices:
        return "A"
    
    # 处理答案格式
    answer_str = str(answer).strip().upper()
    if len(answer_str) > 0 and answer_str[0] in valid_choices:
        return answer_str[0]
    return valid_choices[0]  # 默认返回第一个选项

def get_default_question(course_content: str) -> Dict:
    """生成默认问题（备用）"""
    return {
        "type": "choice",
        "content": json.dumps({
            "question": f"关于{course_content[:50]}...的基本概念是什么？",
            "options": ["A. 基础概念", "B. 进阶内容", "C. 其他"]
        }, ensure_ascii=False),
        "correct_answer": "A",
        "difficulty": 3,
        "timing": "pre_class"
    }
    



# 将习题和问卷生成word文档

# def save_to_word(content, filename="output.docx"):
#     # 解析 Markdown 为纯文本（去掉 #，但保留结构）
#     md_lines = content.split("\n")
#
#     # 创建 Word 文档
#     doc = Document()
#
#     for line in md_lines:
#         if line.startswith("# "):  # 一级标题
#             doc.add_heading(line[2:], level=1)
#         elif line.startswith("## "):  # 二级标题
#             doc.add_heading(line[3:], level=2)
#         elif line.startswith("### "):  # 三级标题
#             doc.add_heading(line[4:], level=3)
#         else:
#             doc.add_paragraph(line)
#
#     # 保存 Word 文件
#     doc.save(filename)
#     print(f"\n✅ Word 文件已保存为：{filename}")



# 生成结构化教案（按六大模块分段）
def generate_lesson_plans(course_content, student_feedback, db_names, similarity_threshold, chunk_cnt):
    """
    根据课程内容、学生反馈和知识库检索结果，生成教学方案
    :param course_content: 课程内容
    :param student_feedback: 学生反馈
    :param db_names: 知识库名称列表
    :param similarity_threshold: 检索相似度阈值
    :param chunk_cnt: 检索片段数量
    :param data_type_filter: 数据类型过滤条件
    :return: 教案内容（Markdown格式）
    """
    # 从知识库检索相关内容
    model_context, display_chunks, source_dict = _retrieve_chunks_from_multiple_dbs(
        query=course_content, db_names=db_names, similarity_threshold=similarity_threshold, chunk_cnt=chunk_cnt
    )
    print("检索内容：")
    print(model_context)
    response = client.chat.completions.create(
        model="deepseek-reasoner",
        messages=[
            {
                "role": "system",
                "content": (
                    "你是教学设计专家，请根据以下三部分信息生成教案：\n"
                    "1. 教师提供的课程内容\n"
                    "2. 学生答题反馈\n"
                    "3. 从知识库检索到的相关教学参考资料（含权威来源）\n\n"
                    "请严格按照以下六个模块生成教案，并满足要求：\n"
                    "### 1. 教学目标\n"
                    "- 列出3-5个可衡量的学习目标，参考知识库中的课程标准。\n\n"
                    "### 2. 教学重难点\n"
                    "- 结合学生反馈和知识库内容，说明重点与难点及突破策略。\n\n"
                    "### 3. 教学内容\n"
                    "- 使用知识库中的资料补充知识点逻辑关系，并用 Mermaid 生成结构图。\n\n"
                    "### 4. 教学时间安排\n"
                    "- 45分钟课时分配，需包含知识库推荐的互动时间比例。\n\n"
                    "### 5. 教学过程\n"
                    "- 按“导入-讲授-互动-小结”设计，互动环节必须包含：\n"
                    "  a) 知识库推荐的活动（如小组讨论/角色扮演）\n"
                    "  b) 针对学生反馈的薄弱点设计练习\n"
                    "- 每个环节需说明：方法、师生行为、时间、工具、预期成果。\n\n"
                    "### 6. 课后作业\n"
                    "- 基础题（覆盖知识库核心内容）\n"
                    "- 拓展题（结合检索到的拓展资料）\n\n"
                    "### 其他要求\n"
                    "- 使用 Markdown 格式，字数不少于3000字\n"
                    "- 关键教学策略需标注来源（如：\"根据[知识库]建议...\"）"
                )
            },
            {
                "role": "user",
                "content": (
                    f"### 课程原始内容\n{course_content}\n\n"
                    f"### 学生反馈\n{student_feedback}\n\n"
                    f"### 知识库参考资料\n{model_context}\n\n"
                    "请生成教案，确保整合上述所有信息。"
                )
            }
        ],
        temperature=1.0,
        max_tokens=32768,  
        top_p=0.9,
        frequency_penalty=0.2,
        presence_penalty=0.1,
        stream=False
    )

    return  response.choices[0].message.content
    



system_prompt = """你是教学设计专家，请根据知识点内容生成至少一道课后练习题。
要求：
1. 返回格式为包含 questions 字段的 JSON 对象，questions 值为题目列表，每个题目包含：
   - type: 题目类型(choice/fill/short_answer)
   - content: 题目内容（选择题需包含question和options字段）
   - correct_answer: 正确答案
   - difficulty: 难度等级(1-5)
   - knowledge_point_id: 知识点ID
2. 必须生成两道题目，一道简单（难度1-2），另一道中等或困难（难度3-5）
3. 题目类型要多样，两道题不能都是同一种类型
4. 用中文回答
5. 确保 knowledge_point_id 与提供的知识点ID一致

示例格式:
{
  "questions": [
    {
      "type": "choice",
      "content": {
        "question": "Python中如何定义变量？",
        "options": ["A. var name = value", "B. name = value", "C. set name value"]
      },
      "correct_answer": "B",
      "difficulty": 2,
      "knowledge_point_id": 1
    },
    {
      "type": "fill",
      "content": "Python中输出文本到控制台的函数是____",
      "correct_answer": "print()",
      "difficulty": 3,
      "knowledge_point_id": 1
    }
  ]
}
"""
def generate_post_class_questions(lesson_plan_content: str, mind_map: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    调用 AI 接口根据教案内容和思维导图生成课后习题，返回符合 Question 数据模型的题目列表。
    每个思维导图的知识叶子节点相关题目不少于两道，且难度不等。
    """
    
    valid_questions = []  # 存储所有有效题目
    processed_node_ids = set()  # 缓存已处理的节点ID

    try:
        # 收集思维导图中的所有叶子节点
        leaf_nodes = []
        
        # 确保 mind_map 是一个列表且包含节点
        if isinstance(mind_map, list) and len(mind_map) > 0 and "data" in mind_map[0]:
            collect_leaf_nodes(mind_map[0], leaf_nodes)
        else:
            collect_leaf_nodes(mind_map, leaf_nodes)
        
        # 如果没有叶子节点，直接返回空列表
        if not leaf_nodes:
            return []

        batch_size = 3  # 每批处理的节点数
        futures = []  # 存储线程池任务

        # 使用线程池并行生成题目
        with ThreadPoolExecutor(max_workers=5) as executor:
            # 提交所有任务到线程池
            for i in range(0, len(leaf_nodes), batch_size):
                batch_leaf_nodes = leaf_nodes[i:i + batch_size]
                futures.append(executor.submit(generate_questions_for_batch, batch_leaf_nodes, lesson_plan_content, processed_node_ids))

            # 收集所有结果
            for future in as_completed(futures):
                batch_questions = future.result()
                valid_questions.extend(batch_questions)

        # 最终校验每个知识点的题目数量
        knowledge_counts = {}
        for q in valid_questions:
            knowledge_counts[q["knowledge_point_id"]] = knowledge_counts.get(q["knowledge_point_id"], 0) + 1

        # 检查并提示题目缺失情况
        for node in leaf_nodes:
            if knowledge_counts.get(node["id"], 0) < 2:
                print(f"警告：知识点 {node['id']} 最终题目数为 {knowledge_counts.get(node['id'], 0)}")

        return valid_questions

    except Exception as e:
        print(f"整体流程错误: {str(e)[:200]}")
        return []


def collect_leaf_nodes(node: Dict[str, Any], leaf_nodes: List[Dict[str, Any]]):
    """递归收集思维导图中的所有叶子节点"""
    # 检查当前节点是否有子节点
    if "children" not in node or len(node.get("children", [])) == 0:
        # 如果没有子节点，且当前节点的 "data" 字段存在，则认为是叶子节点
        if "data" in node and "id" in node["data"] and "text" in node["data"]:
            leaf_nodes.append({
                "id": node["data"]["id"],
                "content": node["data"]["text"],
                "note": node["data"].get("note", ""),
                "depth": node.get("depth", 0)
            })
    else:
        # 递归处理子节点
        for child in node.get("children", []):
            collect_leaf_nodes(child, leaf_nodes)


def generate_questions_for_batch(batch_leaf_nodes: List[Dict[str, Any]], lesson_plan_content: str, processed_node_ids: set) -> List[Dict[str, Any]]:
    """为一批叶子节点生成题目"""
    batch_questions = []

    for node in batch_leaf_nodes:
        if node['id'] in processed_node_ids:
            continue

        processed_node_ids.add(node['id'])
        print(node)
        try:
            # 构造用户消息
            user_message = f"""知识点ID：{node['id']}
知识点内容（标题）：
{node['content']}

知识点内容（备注）：
{node['note']}

教案相关内容（供参考）：
{lesson_plan_content}

请为该知识点生成两道不同难度、不同类型的题目："""

            # 调用 AI 接口生成题目
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ],
                response_format={"type": "json_object"},
                temperature=0.7
            )

            # 解析 AI 响应
            ai_response = json.loads(response.choices[0].message.content)
            questions_data = ai_response.get("questions", [])

            # 题目格式转换和校验
            for q in questions_data:
                # 必要字段校验
                required_fields = ["type", "content", "correct_answer", "difficulty", "knowledge_point_id"]
                if not all(key in q for key in required_fields):
                    continue

                # 知识点ID校验
                if q["knowledge_point_id"] != node["id"]:
                    continue

                # 构造题目对象
                question = {
                    "type": q["type"],
                    "difficulty": min(max(int(q["difficulty"]), 1), 5),
                    "timing": "post_class",
                    "knowledge_point_id": q["knowledge_point_id"]
                }

                # 根据题型处理内容
                if q["type"] == "choice":
                    if isinstance(q["content"], dict) and "question" in q["content"] and "options" in q["content"]:
                        question["content"] = json.dumps(q["content"], ensure_ascii=False)
                        question["correct_answer"] = str(q["correct_answer"]).strip().upper()
                    else:
                        continue  # 跳过格式错误的选择题
                else:
                    question["content"] = str(q["content"])
                    question["correct_answer"] = str(q["correct_answer"])

                batch_questions.append(question)

        except Exception as e:
            print(f"处理知识点 {node['id']} 时出错: {str(e)[:200]}")

    return batch_questions
    



# 教案保存
def save_to_markdown(content, filename="教案.md"):
    """保存教案为 Markdown 文件"""
    with open(filename, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"\n✅ 教案已保存为 Markdown 文件：{filename}")



















# 第一轮：精准知识点提取提示词 (优化后)
extraction_prompt = """请严格遵循以下要求从教学内容中提取纯技术知识点：
## 核心要求
1. 仅提取技术性内容（概念/原理/方法/工具）
2. 排除所有教学相关元素：
   - 教学步骤/时间安排/互动活动
   - 练习题/考试要求/学习目标
   - 师生对话/示例讲解/学习建议
3. 保持技术术语的原始表达：
   - 不修改专业术语
   - 不添加解释性文字
   - 不简化复杂表述

## 结构化要求
4. 按逻辑层级组织：
   - 核心概念 → 子概念 → 技术细节
   - 使用Markdown层级格式：
     ```
     # 主知识点
     ## 子知识点
     ### 技术细节
     ```
5. 复杂概念拆解：
   - 对复合概念进行原子化拆分
   - 保持最小知识单元完整性

待处理内容：
"""

# 第二轮：结构化生成提示词 (优化后)
generation_prompt = """请生成满足以下严格要求的思维导图JSON：

## 数据结构规范
{
  "name": "根节点名称",
  "content": "可选说明",
  "children": [
    {
      "name": "子节点",
      "content": "技术细节",
      // 最多6层深度
    }
  ]
}

## 内容要求
1. 节点命名规范：
   - 名称≤5个汉字或12个英文单词
   - 使用标准技术术语
2. 技术细节要求：
   - 使用Markdown格式：`代码`、**重点**、列表项
   - 包含关键参数/公式/接口等核心信息
3. 逻辑完整性：
   - 保持技术体系的依赖关系
   - 同级节点按技术演进顺序排列
4. 深度控制：
   - 超过6层的分支进行技术聚合
   - 末端节点必须包含实质内容

## 禁止行为
! 禁止创建空内容节点
! 禁止修改原始技术术语

待结构化知识点：
"""

def generate_knowledge_mind_map(lesson_content: str) -> Dict:
    """
    通过严格类型检查的多轮对话生成知识点思维导图
    
    参数：
        lesson_content: 教学内容文本（自动处理为字符串）
        
    返回：
        JSON格式的思维导图数据
    """
    lesson_content = str(lesson_content)
    
    try:
        # 第一阶段：知识点提取
        phase1_response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "你是精准的知识提取引擎"},
                {"role": "user", "content": extraction_prompt + lesson_content}
            ],
            temperature=0.1
        )
        
        # 第二阶段：结构生成
        knowledge_data = phase1_response.choices[0].message.content
        
        phase2_response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "你是知识图谱架构师"},
                {"role": "user", "content": generation_prompt + knowledge_data}
            ],
            response_format={"type": "json_object"},
            temperature=0
        )
        
        raw_response = phase2_response.choices[0].message.content
        processed = preprocess_response(raw_response)
        return validate_structure(processed)
        
    except Exception as e:
        print(f"生成过程中发生错误: {str(e)}")
        return get_default_structure()

def preprocess_response(raw: Any) -> Dict:
    """响应数据标准化处理"""
    if isinstance(raw, str):
        try:
            data = json.loads(raw)
        except:
            data = {"name": "解析失败", "children": []}
    else:
        data = raw

    # 结构修正逻辑
    if "name" not in data:
        return {
            "name": "知识体系",
            "children": [convert_node(k, v) for k, v in data.items()]
        }
    return data

def convert_node(key: str, value: Any) -> Dict:
    """递归转换键值对为节点"""
    node = {"name": key}
    if isinstance(value, dict):
        node["children"] = [convert_node(k, v) for k, v in value.items()]
    elif isinstance(value, list):
        node["children"] = [{"name": str(item)} for item in value]
    else:
        node["content"] = str(value)
    return node

def validate_structure(data: Dict) -> Dict:
    """验证并补全思维导图结构"""
    if not isinstance(data, dict):
        return get_default_structure()
    
    # 确保必填字段存在
    if "name" not in data:
        data["name"] = "知识体系"
    
    # 确保children是列表
    if "children" not in data:
        data["children"] = []
    elif not isinstance(data["children"], list):
        data["children"] = []
    
    return data

def get_default_structure() -> Dict:
    """返回安全的默认结构"""
    return {
        "name": "知识点思维导图",
        "children": []
    }

    
if __name__ == "__main__":
    # 模拟教案内容
    lesson_plan_content = """
    本次课程主要介绍Python的基本语法和数据类型。包括变量的定义、数据类型（如整数、浮点数、字符串、列表、元组、字典等）、条件语句和循环语句。通过案例讲解和互动练习，帮助学生掌握Python编程的基础知识。
    """

    # 模拟思维导图数据（JSON格式）
    mind_map = {"id": 22, "name": "TCP\u8fde\u63a5\u539f\u7406\u6280\u672f\u77e5\u8bc6\u70b9", "content": "", "is_leaf": false, "children": [{"id": 23, "name": "TCP\u534f\u8bae\u57fa\u7840", "content": "", "is_leaf": false, "children": [{"id": 24, "name": "TCP\u5b9a\u4e49\u4e0e\u7279\u6027", "content": "- **\u9762\u5411\u8fde\u63a5\u7684\u4f20\u8f93\u5c42\u534f\u8bae**\n- **\u53ef\u9760\u6027\u4f20\u8f93\u673a\u5236**\n- **\u6d41\u91cf\u63a7\u5236\u529f\u80fd**\n- **\u62e5\u585e\u63a7\u5236\u529f\u80fd**", "is_leaf": true, "children": []}, {"id": 25, "name": "TCP\u4e0eUDP\u5bf9\u6bd4", "content": "- **TCP\u7279\u6027**\uff1a\u53ef\u9760\u6027\u3001\u6709\u5e8f\u6027\u3001\u9762\u5411\u8fde\u63a5\n- **UDP\u7279\u6027**\uff1a\u9ad8\u6548\u7387\u3001\u65e0\u8fde\u63a5\u3001\u65e0\u5e8f\u6027", "is_leaf": true, "children": []}, {"id": 26, "name": "\u534f\u8bae\u6808\u4f4d\u7f6e", "content": "- **\u4f20\u8f93\u5c42\u534f\u8bae**\n- **\u5de5\u4f5c\u5728IP\u534f\u8bae\u4e4b\u4e0a**", "is_leaf": true, "children": []}]}, {"id": 27, "name": "TCP\u8fde\u63a5\u7ba1\u7406", "content": "", "is_leaf": false, "children": [{"id": 28, "name": "\u4e09\u6b21\u63e1\u624b\u5efa\u7acb\u8fde\u63a5", "content": "- **SYN\u62a5\u6587**\uff1a\u521d\u59cb\u5e8f\u5217\u53f7\u4ea4\u6362\n- **SYN-ACK\u62a5\u6587**\uff1a\u786e\u8ba4\u5e8f\u5217\u53f7+\u670d\u52a1\u5668\u521d\u59cb\u5e8f\u5217\u53f7\n- **ACK\u62a5\u6587**\uff1a\u786e\u8ba4\u670d\u52a1\u5668\u5e8f\u5217\u53f7\n- **\u5e8f\u5217\u53f7(Sequence Number)\u4f5c\u7528**\uff1a\u6807\u8bc6\u6570\u636e\u5b57\u8282\u6d41\n- **\u786e\u8ba4\u53f7(Acknowledgment Number)\u4f5c\u7528**\uff1a\u671f\u671b\u63a5\u6536\u7684\u4e0b\u4e00\u4e2a\u5b57\u8282\u5e8f\u53f7\n- **\u8fde\u63a5\u72b6\u6001**\uff1a\n  - \u534a\u8fde\u63a5\u72b6\u6001(SYN_RECEIVED)\n  - \u5168\u8fde\u63a5\u72b6\u6001(ESTABLISHED)", "is_leaf": true, "children": []}, {"id": 29, "name": "\u56db\u6b21\u6325\u624b\u65ad\u5f00\u8fde\u63a5", "content": "- **FIN\u62a5\u6587**\uff1a\u53d1\u8d77\u8fde\u63a5\u7ec8\u6b62\n- **ACK\u62a5\u6587**\uff1a\u786e\u8ba4\u7ec8\u6b62\u8bf7\u6c42\n- **\u72b6\u6001\u8f6c\u6362**\uff1a\n  - FIN_WAIT_1\n  - FIN_WAIT_2\n  - TIME_WAIT(\u7b49\u5f852MSL\u65f6\u95f4)\n- **TIME_WAIT\u72b6\u6001\u610f\u4e49**\uff1a\n  - \u786e\u4fdd\u6700\u540e\u4e00\u4e2aACK\u5230\u8fbe\n  - \u8ba9\u7f51\u7edc\u4e2d\u6b8b\u7559\u62a5\u6587\u6bb5\u5931\u6548", "is_leaf": true, "children": []}]}, {"id": 30, "name": "TCP\u6570\u636e\u4f20\u8f93\u673a\u5236", "content": "", "is_leaf": false, "children": [{"id": 31, "name": "\u53ef\u9760\u6027\u4fdd\u8bc1\u673a\u5236", "content": "- **\u786e\u8ba4\u5e94\u7b54(ACK)**\uff1a\u63a5\u6536\u65b9\u53d1\u9001\u786e\u8ba4\u62a5\u6587\n- **\u8d85\u65f6\u91cd\u4f20**\uff1a\u672a\u6536\u5230ACK\u65f6\u91cd\u53d1\u6570\u636e\n- **\u6570\u636e\u6392\u5e8f**\uff1a\u901a\u8fc7\u5e8f\u5217\u53f7\u91cd\u7ec4\u4e71\u5e8f\u62a5\u6587", "is_leaf": true, "children": []}, {"id": 32, "name": "\u6d41\u91cf\u63a7\u5236", "content": "- **\u6ed1\u52a8\u7a97\u53e3\u539f\u7406**\uff1a\n  - \u63a5\u6536\u7a97\u53e3(rwnd)\uff1a\u63a5\u6536\u65b9\u7f13\u51b2\u533a\u5927\u5c0f\n  - \u53d1\u9001\u7a97\u53e3\uff1a\u4e0d\u8d85\u8fc7\u63a5\u6536\u7a97\u53e3\u548c\u62e5\u585e\u7a97\u53e3\n- **\u7a97\u53e3\u901a\u544a**\uff1a\u901a\u8fc7TCP\u5934\u90e8\u7a97\u53e3\u5b57\u6bb5\u901a\u77e5\u5bf9\u7aef", "is_leaf": true, "children": []}, {"id": 33, "name": "\u62e5\u585e\u63a7\u5236", "content": "- **\u6162\u542f\u52a8\u7b97\u6cd5**\uff1a\u7a97\u53e3\u5927\u5c0f\u6307\u6570\u589e\u957f\n- **\u62e5\u585e\u907f\u514d\u7b97\u6cd5**\uff1a\u7a97\u53e3\u5927\u5c0f\u7ebf\u6027\u589e\u957f\n- **\u5feb\u901f\u91cd\u4f20**\uff1a\u6536\u52303\u4e2a\u91cd\u590dACK\u7acb\u5373\u91cd\u4f20\n- **\u5feb\u901f\u6062\u590d**\uff1a\u5feb\u901f\u91cd\u4f20\u540e\u7684\u7a97\u53e3\u8c03\u6574\u7b56\u7565", "is_leaf": true, "children": []}]}, {"id": 34, "name": "TCP\u534f\u8bae\u5b9e\u73b0\u7ec6\u8282", "content": "", "is_leaf": false, "children": [{"id": 35, "name": "\u5173\u952e\u62a5\u6587\u5b57\u6bb5", "content": "- **\u5e8f\u5217\u53f7(32\u4f4d)**\n- **\u786e\u8ba4\u53f7(32\u4f4d)**\n- **\u7a97\u53e3\u5927\u5c0f(16\u4f4d)**\n- **\u6807\u5fd7\u4f4d(SYN/ACK/FIN/RST\u7b49)**", "is_leaf": true, "children": []}, {"id": 36, "name": "\u62e5\u585e\u63a7\u5236\u7b97\u6cd5\u53d8\u79cd", "content": "- **TCP Tahoe**\n- **TCP Reno**", "is_leaf": true, "children": []}, {"id": 37, "name": "\u6027\u80fd\u4f18\u5316\u53c2\u6570", "content": "- **\u6700\u5927\u62a5\u6587\u6bb5\u957f\u5ea6(MSS)**\n- **\u7a97\u53e3\u7f29\u653e\u56e0\u5b50(Window Scaling)**\n- **\u65f6\u95f4\u6233\u9009\u9879(Timestamp)**", "is_leaf": true, "children": []}]}]}

    # 调用生成课后习题的函数
    generated_questions = generate_post_class_questions(lesson_plan_content, mind_map)

    # 打印生成的课后习题
    print("生成的课后习题如下：")
    print(json.dumps(generated_questions, ensure_ascii=False, indent=2))

from app.config import Config

# 题型中文映射
q_type_dict = {
    "choice": "选择题",
    "fill": "填空题",
    "short_answer": "简答题"
}

def generate_questions_with_specs(
    knowledge_point_name: str,
    knowledge_point_content: str,
    question_specs: List[Tuple[str, int, Optional[int]]],
    retrieved_content: str = ""
) -> List[Dict[str, any]]:
    """
    使用DeepSeek-Reasoner模型生成指定类型、数量和难度的题目
    
    参数:
    - knowledge_point_name: 知识点名称
    - knowledge_point_content: 知识点详细内容
    - question_specs: 题目规格列表，每个元素为(题型, 数量, 难度)
        - 题型支持: 'choice', 'fill', 'short_answer'
        - 难度: 1-5级 (1最简单，5最难)，None表示随机难度
    - retrieved_content: 从知识库检索到的相关内容（可选）
    
    返回:
    - 题目字典列表，格式为：
        [{
            "type": 题目类型,
            "content": 题目内容,
            "correct_answer": 正确答案,
            "difficulty": 难度等级(1-5),
            "knowledge_point": 关联知识点
        }]
    """
    
    # 验证输入规格
    valid_types = {'choice', 'fill', 'short_answer'}
    for spec in question_specs:
        if len(spec) != 3:
            raise ValueError("每个题目规格必须是(题型, 数量, 难度)三元组")
        if spec[0] not in valid_types:
            raise ValueError(f"无效题型: {spec[0]}，支持类型: {', '.join(valid_types)}")
        if spec[1] < 1:
            raise ValueError("题目数量必须至少为1")
        if spec[2] is not None and (spec[2] < 1 or spec[2] > 5):
            raise ValueError("难度必须在1-5之间")
    
    # 构建题型要求描述
    type_requirements = []
    for i, (q_type, count, difficulty) in enumerate(question_specs):
        diff_desc = f"难度{difficulty}" if difficulty else "随机难度"
        type_requirements.append(f"{i+1}. {count}道{q_type_dict[q_type]} ({diff_desc})")
    
    # 构建系统提示词
    system_prompt = f"""
    你是一位资深教师，需要根据知识点和参考资料生成课后习题。请严格遵守以下要求：
    1. 严格按以下规格生成题目：
        {chr(10).join(type_requirements)}
    2. 所有题目必须关联知识点：{knowledge_point_name}
    3. 选择题需提供4个选项和正确答案字母
    4. 使用以下JSON格式返回：
        {{"questions": [
            {{
                "type": "题型标识",(题型仅支持: 'choice', 'fill', 'short_answer')
                "content": "题目内容",
                "correct_answer": "正确答案",
                "difficulty": 难度等级,
                "analysis": "题目解析"
            }}
        ]}}
    5. 确保题目清晰、答案准确、难度符合要求
    6. 题目内容严格按照markdown格式输出
    7.  其中的content字段的内容必须严格按照markdown格式进行输出，如果引用参考资源，则保持参考资料中的markdown格式，同时保留图片
    """
    
    # 构建用户输入
    user_input = f"""
    ## 核心知识点
    名称：{knowledge_point_name}
    内容：{knowledge_point_content}
    
    ## 参考资料
    {retrieved_content if retrieved_content else "无相关参考资料"}
    
    ## 生成要求
    请严格按照上述规格生成题目，确保题型、数量和难度要求准确无误。
    """
    
    try:
        # 调用DeepSeek-Reasoner API
        response = requests.post(
            "https://api.deepseek.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {Config.DEEPSEEK_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": "deepseek-reasoner",
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_input}
                ],
                "response_format": {"type": "json_object"},
                "temperature": 0.7,
                "max_tokens": 3000
            }
        )
        
        # 检查响应状态
        if response.status_code != 200:
            logger.error(f"DeepSeek API请求失败: {response.status_code} - {response.text}")
            return []
        
        # 解析API响应
        response_data = response.json()
        ai_content = response_data["choices"][0]["message"]["content"]
        print(ai_content)
        # 尝试解析JSON
        try:
            questions_data = json.loads(ai_content)["questions"]
        except (json.JSONDecodeError, KeyError) as e:
            logger.error(f"响应解析失败: {str(e)}")
            logger.debug(f"原始响应内容: {ai_content[:500]}...")
            return []
        print(questions_data)
        # 格式化题目内容
        formatted_questions = []
        for q in questions_data:

            # 确保难度在1-5范围内
            if "difficulty" in q:
                q["difficulty"] = max(1, min(5, int(q["difficulty"])))
            else:
                q["difficulty"] = 3  # 默认难度
            
            formatted_questions.append(q)
        
        return formatted_questions
    
    except requests.Timeout:
        logger.error("DeepSeek API请求超时")
        return []
    except Exception as e:
        logger.error(f"题目生成失败: {str(e)}")
        return []